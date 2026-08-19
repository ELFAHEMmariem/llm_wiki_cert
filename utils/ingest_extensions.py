# utils/ingest_extensions.py
import re
import docx

def process_docx(file_path: str) -> str:
    """Extrait le texte structuré d'un fichier Microsoft Word (.docx)."""
    doc = docx.Document(file_path)
    full_text = []
    
    for para in doc.paragraphs:
        if para.text.strip():
            full_text.append(para.text.strip())
            
    # Extraction simple des tableaux Word
    for table in doc.tables:
        for row in table.rows:
            row_data = [cell.text.strip() for cell in row.cells]
            full_text.append(" | ".join(row_data))
            
    return "\n\n".join(full_text)


def process_markdown_clipper(file_path: str) -> str:
    """Traitement des fichiers Markdown natifs ou issus d'Obsidian Web Clipper."""
    with open(file_path, "r", encoding="utf-8") as f:
        content = f.read()

    # Nettoyage léger du YAML Frontmatter s'il existe (--- ... ---)
    content = re.sub(r"^---\n.*?---\n", "", content, flags=re.DOTALL)
    
    return content.strip()